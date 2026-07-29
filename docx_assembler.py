#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
특기시방서 Word 문서 조립기 (python-docx 버전)
"""

import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_font(run, name='맑은 고딕', size=10, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if color:
        run.font.color.rgb = RGBColor(*color)


def _heading_para(doc, text, level=1, size=14, color=(31, 78, 121)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _set_font(run, size=size, bold=True, color=color)
    return p


def _body_para(doc, text, size=10.5, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    _set_font(run, size=size)
    return p


def _set_page_margins(section, top=2.5, bottom=2.5, left=3.0, right=2.5):
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


def _add_cover(doc, project_name, subtitle, client, year_month):
    """표지 추가"""
    for _ in range(10):
        doc.add_paragraph()

    # 제목
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(project_name)
    _set_font(run, size=20, bold=True, color=(31, 78, 121))

    # 부제목
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'({subtitle})')
    _set_font(run, size=15, color=(46, 117, 182))

    for _ in range(6):
        doc.add_paragraph()

    # 연월
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(year_month)
    _set_font(run, size=13)

    for _ in range(3):
        doc.add_paragraph()

    # 발주처
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(client)
    _set_font(run, size=14, bold=True)

    doc.add_page_break()


def _add_toc(doc, chapter_names):
    """목차"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('목  차')
    _set_font(run, size=14, bold=True)
    doc.add_paragraph()

    # 공사개요
    p = doc.add_paragraph()
    run = p.add_run('< 공 사 개 요 >')
    _set_font(run, size=11)

    for i, name in enumerate(chapter_names):
        p = doc.add_paragraph()
        run = p.add_run(f'<제 {i+1} 장  {name}>')
        _set_font(run, size=11)

    doc.add_page_break()


def _add_summary(doc, info):
    """공사개요"""
    _heading_para(doc, '공  사  개  요', size=14)
    doc.add_paragraph()

    _body_para(doc, f'1. 공  사  명  :  {info.get("project_name", "")}')
    _body_para(doc, f'2. 발    주    처  :  {info.get("client", "")}')
    _body_para(doc, f'3. 대  지  위  치  :  {info.get("location", "")}')
    doc.add_paragraph()

    # 설계개요 표
    _body_para(doc, '4. 설 계 개 요')
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    design_data = [
        ('대지면적', info.get('site_area',''), '건축면적', info.get('building_area','')),
        ('연  면  적', info.get('total_floor_area',''), '건  폐  율', info.get('building_coverage','')),
        ('용  적  율', info.get('floor_area_ratio',''), '층      수', info.get('floors','')),
        ('주요구조', info.get('structure',''), '기      초', info.get('foundation','')),
    ]
    for i, (k1, v1, k2, v2) in enumerate(design_data):
        row = table.rows[i]
        for j, (text, is_header) in enumerate([(k1,True),(v1,False),(k2,True),(v2,False)]):
            cell = row.cells[j]
            cell.text = text
            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(text)
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].text = text
            _set_font(cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(text),
                      size=10, bold=is_header)
            if is_header:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'D9E1F2')
                tcPr.append(shd)
    doc.add_paragraph()

    # 층별 바닥면적
    floor_areas = info.get('floor_areas', [])
    if floor_areas:
        _body_para(doc, '5. 층별 바닥면적')
        ft = doc.add_table(rows=len(floor_areas)+1, cols=3)
        ft.style = 'Table Grid'
        for j, h in enumerate(['층별', '면적(㎡)', '비고']):
            cell = ft.rows[0].cells[j]
            cell.text = h
            _set_font(cell.paragraphs[0].add_run(h) if not cell.paragraphs[0].runs else cell.paragraphs[0].runs[0],
                      size=10, bold=True)
        for i, row_data in enumerate(floor_areas):
            for j, val in enumerate(row_data[:3]):
                cell = ft.rows[i+1].cells[j]
                cell.text = str(val)
        doc.add_paragraph()

    # 재료마감
    finishes = info.get('finishes', [])
    if finishes:
        _body_para(doc, '6. 재료 마감')
        for k, v in finishes:
            _body_para(doc, f'  {k}: {v}', indent=0.5)

    doc.add_page_break()


def _add_spec_body(doc, specs):
    """시방서 본문"""
    for i, spec in enumerate(specs):
        if spec.get('error'):
            continue

        # 장 제목
        _heading_para(doc, f'제 {i+1} 장  {spec["input_name"]}', size=13)

        # KCS 코드 표시
        if spec.get('detail_code'):
            status_text = {'STANDARD':'표준시방서 원문 인용','STANDARD_REF':'근접 표준시방서 참조','SPEC_ONLY':'특기시방서(AI 초안)'}.get(spec.get('status',''), '')
            p = doc.add_paragraph()
            run = p.add_run(f'[{status_text}]  {spec["detail_code"]}  {spec.get("detail_name","") or ""}')
            _set_font(run, size=9, color=(89, 89, 89))

        for sec in (spec.get('sections') or []):
            if sec.get('is_ai'):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(8)
                run = p.add_run('⚠  AI 초안 — 검수필요  ⚠')
                _set_font(run, size=10, bold=True, color=(192, 0, 0))

            body_text = sec.get('body', '')
            for line in body_text.split('\n'):
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = Pt(17)
                run = p.add_run(line.rstrip())
                _set_font(run, size=10)

        if i < len(specs) - 1:
            doc.add_page_break()


def assemble_docx(project_info: dict, specs: list, output_dir: str,
                  filename: str = None) -> str:
    if not filename:
        safe = "".join(
            c for c in project_info.get("project_name", "특기시방서")
            if c.isalnum() or '가' <= c <= '힣' or c in " _-"
        ).strip()
        filename = f"{safe or '특기시방서'}_특기시방서.docx"

    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    out_path = output_dir / filename

    doc = Document()
    section = doc.sections[0]
    _set_page_margins(section)

    pname = project_info.get('project_name', '')
    client = project_info.get('client', '')
    ym = project_info.get('year_month', datetime.datetime.now().strftime('%Y. %m.'))

    # 표지 2장
    _add_cover(doc, pname, '건축, 기계시방서', client, ym)
    _add_cover(doc, pname, '건축시방서', client, ym)

    # 목차
    chapter_names = [s['input_name'] for s in specs if not s.get('error')]
    _add_toc(doc, chapter_names)

    # 공사개요
    _add_summary(doc, project_info)

    # 시방서 본문
    _add_spec_body(doc, specs)

    doc.save(str(out_path))
    return str(out_path)


if __name__ == "__main__":
    import datetime
    info = {
        "project_name": "테스트 신축공사", "client": "테스트 발주처",
        "year_month": "2026. 07.", "location": "서울시 강남구",
        "site_area": "500㎡", "building_area": "200㎡",
        "total_floor_area": "400㎡", "building_coverage": "40%",
        "floor_area_ratio": "80%", "floors": "지상 2층",
        "structure": "RC조", "foundation": "줄기초",
        "floor_areas": [["1층","200",""],["2층","200",""],["합계","400",""]],
        "finishes": [["외벽","뿜칠"],["바닥","강화마루"]],
    }
    specs = [{"input_name":"콘크리트공사","status":"STANDARD",
              "detail_code":"KCS 41 30 00","detail_name":"콘크리트공사",
              "sections":[{"is_ai":False,"body":"1. 일반사항\n본 시방에 따라 시공한다."}],
              "error":None}]
    out = assemble_docx(info, specs, "/tmp", "test.docx")
    print(f"생성: {out} ({Path(out).stat().st_size:,} bytes)")
